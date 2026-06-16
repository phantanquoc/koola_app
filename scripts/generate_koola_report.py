from __future__ import annotations

from pathlib import Path
import shutil
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "report-assets"
OUT = DOCS / "Bao_cao_tong_quat_APP_KOOLA.docx"

WIDTH = 1080
HEIGHT = 1920

BLUE = "#2563EB"
GREEN = "#10B981"
WARM = "#F97316"
DANGER = "#EF4444"
INK = "#111827"
MUTED = "#667085"
FAINT = "#98A2B3"
LINE = "#E5E7EB"
CANVAS = "#F6F7FB"
SURFACE = "#FFFFFF"

FONT_DIR = Path("C:/Windows/Fonts")
FONT_REGULAR = FONT_DIR / "arial.ttf"
FONT_BOLD = FONT_DIR / "arialbd.ttf"


def image_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    font_path = FONT_BOLD if bold else FONT_REGULAR
    if font_path.exists():
        return ImageFont.truetype(str(font_path), size)
    return ImageFont.load_default()


def draw_text(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    value: str,
    size: int = 32,
    fill: str = INK,
    bold: bool = False,
    anchor: str | None = None,
    max_width: int | None = None,
) -> None:
    font = image_font(size, bold=bold)
    if max_width is None:
        draw.text(xy, value, font=font, fill=fill, anchor=anchor)
        return

    words = value.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if not current or draw.textlength(candidate, font=font) <= max_width:
            current = candidate
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)

    x, y = xy
    line_height = size + 8
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        y += line_height


def rounded(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    radius: int,
    fill: str,
    outline: str | None = None,
    width: int = 1,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def status_bar(draw: ImageDraw.ImageDraw) -> None:
    draw_text(draw, (72, 42), "3:10", size=34, fill="#5E6470")
    draw.polygon([(900, 56), (930, 56), (930, 86)], fill="#5E6470")
    draw.rectangle((966, 50, 1010, 86), fill="#5E6470")
    draw.rectangle((1014, 60, 1020, 76), fill="#5E6470")


def logo(draw: ImageDraw.ImageDraw, x: int, y: int, size: int = 58) -> None:
    cursor = x
    for letter, color in [("K", BLUE), ("O", GREEN), ("O", WARM), ("L", BLUE), ("A", GREEN)]:
        draw_text(draw, (cursor, y), letter, size=size, fill=color, bold=True)
        cursor += int(size * 0.68)


def bottom_tabs(draw: ImageDraw.ImageDraw, active: str) -> None:
    y = HEIGHT - 154
    draw.rectangle((0, y, WIDTH, HEIGHT), fill=SURFACE)
    draw.line((0, y, WIDTH, y), fill=LINE, width=2)
    tabs = [
        ("Trò chuyện", "chat"),
        ("Mua sắm", "shopping"),
        ("Kết nối", "connect"),
        ("Dịch vụ", "services"),
        ("Cá nhân", "personal"),
    ]
    slot = WIDTH / len(tabs)
    for index, (label, key) in enumerate(tabs):
        cx = int(slot * index + slot / 2)
        color = BLUE if key == active else "#8C8C8C"
        draw_text(draw, (cx, y + 42), "●", size=42, fill=color, bold=True, anchor="mm")
        draw_text(draw, (cx, y + 92), label, size=26, fill=color, bold=True, anchor="mm")
    draw.rounded_rectangle((395, HEIGHT - 20, 685, HEIGHT - 12), radius=8, fill="#111111")


def phone_canvas() -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (WIDTH, HEIGHT), CANVAS)
    draw = ImageDraw.Draw(image)
    status_bar(draw)
    return image, draw


def chip(draw: ImageDraw.ImageDraw, x: int, y: int, label: str, selected: bool = False) -> int:
    font = image_font(26, bold=True)
    width = int(draw.textlength(label, font=font) + 48)
    fill = BLUE if selected else SURFACE
    color = SURFACE if selected else MUTED
    rounded(draw, (x, y, x + width, y + 52), 26, fill, BLUE if selected else LINE)
    draw_text(draw, (x + 24, y + 12), label, size=26, fill=color, bold=True)
    return x + width + 14


def header(draw: ImageDraw.ImageDraw, search: str, title: str | None = None) -> None:
    draw.rectangle((0, 112, WIDTH, 350), fill=SURFACE)
    if title:
        draw_text(draw, (44, 142), title, size=44, fill=INK, bold=True)
    else:
        logo(draw, 44, 132, 62)
    rounded(draw, (44, 244, 760, 316), 36, "#F3F6FB")
    draw_text(draw, (82, 282), "Q", size=34, fill=MUTED, bold=True, anchor="mm")
    draw_text(draw, (122, 262), search, size=31, fill=MUTED)
    rounded(draw, (798, 240, 892, 334), 47, "#F3F6FB")
    draw_text(draw, (845, 288), "QR", size=23, fill=BLUE, bold=True, anchor="mm")
    rounded(draw, (920, 240, 1014, 334), 47, BLUE)
    draw_text(draw, (967, 287), "+", size=54, fill=SURFACE, bold=True, anchor="mm")
    draw.line((0, 350, WIDTH, 350), fill=LINE, width=2)


def save_chat_asset() -> Path:
    source = ROOT / ".screenshots" / "koola.png"
    target = ASSETS / "tab_tro_chuyen.png"
    if source.exists():
        shutil.copyfile(source, target)
        return target

    image, draw = phone_canvas()
    header(draw, "Tìm kiếm...")
    draw.rectangle((0, 350, WIDTH, 470), fill=SURFACE)
    for index, label in enumerate(["Tin nhắn", "Cuộc gọi", "Danh bạ"]):
        cx = int(WIDTH / 3 * index + WIDTH / 6)
        color = BLUE if index == 0 else FAINT
        draw_text(draw, (cx, 405), label, size=28, fill=color, bold=True, anchor="mm")
    draw.rounded_rectangle((70, 452, 146, 460), radius=4, fill=BLUE)
    rows = [
        ("Quoc", "hi", "3 ngày trước", GREEN, "Q"),
        ("Test Group 5.4", "kk", "11 ngày trước", "#0EA5E9", "T"),
        ("112233", "alo", "13 ngày trước", GREEN, "1"),
    ]
    y = 540
    for name, preview, when, color, initial in rows:
        draw.line((0, y - 28, WIDTH, y - 28), fill=LINE, width=1)
        draw.ellipse((48, y, 174, y + 126), fill=color)
        draw_text(draw, (111, y + 65), initial, size=46, fill=SURFACE, bold=True, anchor="mm")
        draw_text(draw, (216, y + 18), name, size=38, fill=INK, bold=True)
        draw_text(draw, (216, y + 76), preview, size=34, fill=MUTED)
        draw_text(draw, (1032, y + 32), when, size=27, fill=FAINT, bold=True, anchor="ra")
        y += 178
    bottom_tabs(draw, "chat")
    image.save(target)
    return target


def make_shopping_asset() -> Path:
    image, draw = phone_canvas()
    draw.rectangle((0, 112, WIDTH, 224), fill=SURFACE)
    logo(draw, 36, 142, 42)
    rounded(draw, (236, 142, 878, 194), 26, CANVAS)
    draw_text(draw, (300, 150), "Tìm sản phẩm, cửa hàng...", size=24, fill=MUTED)
    rounded(draw, (914, 136, 984, 206), 22, BLUE)
    draw_text(draw, (949, 171), "GH", size=22, fill=SURFACE, bold=True, anchor="mm")
    draw.line((0, 224, WIDTH, 224), fill=LINE, width=2)

    rounded(draw, (44, 258, 1036, 362), 14, SURFACE, LINE)
    rounded(draw, (74, 284, 130, 340), 14, "#FFF4E5")
    draw_text(draw, (102, 313), "%", size=34, fill=WARM, bold=True, anchor="mm")
    draw_text(draw, (156, 280), "Deal nhanh quanh bạn", size=30, bold=True)
    draw_text(draw, (156, 318), "Mua tạp hóa, đồ ăn và vật dụng giao trong ngày", size=23, fill=MUTED)
    draw_text(draw, (970, 288), "-25%", size=30, fill=WARM, bold=True, anchor="ra")

    x = 44
    for label, color in [("Siêu thị", GREEN), ("Ăn uống", WARM), ("Freeship", BLUE), ("Deal sốc", DANGER)]:
        rounded(draw, (x, 388, x + 236, 482), 14, SURFACE, LINE)
        rounded(draw, (x + 82, 404, x + 154, 456), 14, f"{color}22")
        draw_text(draw, (x + 118, 428), "●", size=30, fill=color, bold=True, anchor="mm")
        draw_text(draw, (x + 118, 466), label, size=22, fill=INK, bold=True, anchor="mm")
        x += 254

    draw_text(draw, (44, 530), "Gợi ý mua sắm", size=40, bold=True)
    draw_text(draw, (44, 578), "Sản phẩm nổi bật, cửa hàng gần bạn và deal trong ngày", size=24, fill=MUTED)
    x = 44
    for label, selected in [("Tất cả", True), ("Tạp hóa", False), ("Đồ ăn", False), ("Điện tử", False)]:
        x = chip(draw, x, 626, label, selected)

    products = [
        ("Combo rau củ tươi Đà Lạt", "Koola Fresh Market", "89.000đ", GREEN, "Giao 2h"),
        ("Cơm gà sốt tiêu xanh", "Bếp Nhà Koola", "45.000đ", WARM, "Bán chạy"),
        ("Tai nghe Bluetooth Mini", "Tech Corner", "249.000đ", BLUE, "-22%"),
        ("Bộ lau nhà gấp gọn", "Nhà Xinh Store", "159.000đ", "#14B8A6", "Hot"),
    ]
    for (title, shop, price, color, badge), (x, y) in zip(products, [(44, 700), (558, 700), (44, 1048), (558, 1048)]):
        rounded(draw, (x, y, x + 478, y + 318), 14, SURFACE, LINE)
        rounded(draw, (x, y, x + 478, y + 118), 14, f"{color}22")
        draw_text(draw, (x + 239, y + 64), "SP", size=38, fill=color, bold=True, anchor="mm")
        rounded(draw, (x + 14, y + 14, x + 118, y + 46), 8, INK)
        draw_text(draw, (x + 66, y + 31), badge, size=17, fill=SURFACE, bold=True, anchor="mm")
        draw_text(draw, (x + 18, y + 138), title, size=27, bold=True, max_width=430)
        draw_text(draw, (x + 18, y + 208), shop, size=22, fill=MUTED)
        draw_text(draw, (x + 18, y + 242), "★ 4.8  Đã bán 1.2k", size=22, fill=MUTED)
        draw_text(draw, (x + 18, y + 278), price, size=28, fill=DANGER, bold=True)
        rounded(draw, (x + 410, y + 260, x + 458, y + 308), 12, BLUE)
        draw_text(draw, (x + 434, y + 284), "+", size=30, fill=SURFACE, bold=True, anchor="mm")

    draw_text(draw, (44, 1412), "Cửa hàng gần bạn", size=36, bold=True)
    for index, (name, meta, color) in enumerate([
        ("Koola Mart Nguyễn Huệ", "Tạp hóa · 1.2 km · 18-25 phút", BLUE),
        ("Bếp Cô Ba", "Cơm văn phòng · 800 m · 20 phút", WARM),
    ]):
        y = 1472 + index * 112
        rounded(draw, (44, y, 1036, y + 92), 14, SURFACE, LINE)
        rounded(draw, (70, y + 18, 126, y + 74), 14, f"{color}22")
        draw_text(draw, (98, y + 47), "CH", size=18, fill=color, bold=True, anchor="mm")
        draw_text(draw, (150, y + 15), name, size=28, bold=True)
        draw_text(draw, (150, y + 52), meta, size=22, fill=MUTED)
        draw_text(draw, (1002, y + 46), ">", size=34, fill=FAINT, bold=True, anchor="mm")

    bottom_tabs(draw, "shopping")
    target = ASSETS / "tab_mua_sam.png"
    image.save(target)
    return target


def make_connect_asset() -> Path:
    image, draw = phone_canvas()
    header(draw, "Tìm doanh nghiệp...")
    draw.rectangle((0, 350, WIDTH, 424), fill=SURFACE)
    for index, (label, selected) in enumerate([("Tất cả", True), ("Đối tác", False), ("Nhà cung cấp", False)]):
        cx = int(WIDTH / 3 * index + WIDTH / 6)
        draw_text(draw, (cx, 386), label, size=29, fill=BLUE if selected else MUTED, bold=True, anchor="mm")
        if selected:
            draw.rounded_rectangle((cx - 70, 420, cx + 70, 426), radius=3, fill=BLUE)
    draw.line((0, 424, WIDTH, 424), fill=LINE, width=2)
    x = 44
    for label, selected in [("Tỉnh/Thành", False), ("Mới nhất", False), ("Logistics", True), ("Nội địa", False)]:
        x = chip(draw, x, 454, label, selected)

    rounded(draw, (44, 534, 1036, 646), 14, "#EEF6FF", "#BFDBFE")
    rounded(draw, (70, 562, 126, 618), 14, BLUE)
    draw_text(draw, (98, 590), "DN", size=18, fill=SURFACE, bold=True, anchor="mm")
    draw_text(draw, (150, 558), "Kết nối doanh nghiệp", size=31, bold=True)
    draw_text(draw, (150, 596), "Tìm đối tác, nhà cung cấp và bắt đầu trò chuyện nhanh.", size=23, fill=MUTED)
    rounded(draw, (812, 574, 1008, 616), 12, GREEN)
    draw_text(draw, (910, 596), "Đăng ký", size=23, fill=SURFACE, bold=True, anchor="mm")

    rows = [
        ("Koola Logistics", "Logistics · TP.HCM", "Đối tác vận chuyển nội thành, nhận đơn trong ngày", "Đối tác", BLUE, True),
        ("Bao Bì Việt", "Bao bì · Bình Dương", "Cung cấp thùng carton, túi giấy và in nhãn theo yêu cầu", "Nhà cung cấp", GREEN, False),
        ("FoodTech Supply", "Thực phẩm · Hà Nội", "Nguyên liệu F&B, giao định kỳ cho cửa hàng và bếp trung tâm", "Nhà cung cấp", WARM, True),
        ("Retail Pro", "Bán lẻ · Đà Nẵng", "Kết nối chuỗi cửa hàng địa phương và nhà phân phối", "Đối tác", "#7C3AED", False),
    ]
    y = 684
    for name, category, description, relation, color, verified in rows:
        rounded(draw, (44, y, 1036, y + 184), 16, SURFACE, LINE)
        rounded(draw, (72, y + 28, 154, y + 110), 18, color)
        draw_text(draw, (113, y + 68), "B", size=36, fill=SURFACE, bold=True, anchor="mm")
        verified_suffix = " (xác minh)" if verified else ""
        draw_text(draw, (182, y + 24), f"{name}{verified_suffix}", size=31, fill=INK, bold=True)
        draw_text(draw, (182, y + 62), category, size=23, fill=MUTED)
        draw_text(draw, (182, y + 94), description, size=23, fill=MUTED, max_width=620)
        rounded(draw, (780, y + 28, 1000, y + 68), 20, "#EFF6FF")
        draw_text(draw, (890, y + 49), relation, size=21, fill=BLUE, bold=True, anchor="mm")
        rounded(draw, (182, y + 132, 350, y + 168), 18, BLUE)
        draw_text(draw, (266, y + 151), "Nhắn tin", size=21, fill=SURFACE, bold=True, anchor="mm")
        rounded(draw, (368, y + 132, 540, y + 168), 18, "#ECFDF3")
        draw_text(draw, (454, y + 151), "Kết nối", size=21, fill=GREEN, bold=True, anchor="mm")
        y += 204

    bottom_tabs(draw, "connect")
    target = ASSETS / "tab_ket_noi.png"
    image.save(target)
    return target


def make_services_asset() -> Path:
    image, draw = phone_canvas()
    draw.rectangle((0, 112, WIDTH, 224), fill=SURFACE)
    draw_text(draw, (44, 132), "Dịch vụ", size=42, fill=INK, bold=True)
    draw_text(draw, (44, 178), "Đặt nhanh các nhu cầu hằng ngày", size=22, fill=MUTED)
    rounded(draw, (340, 142, 878, 194), 26, CANVAS)
    draw_text(draw, (410, 150), "Tìm sửa chữa, giao hàng...", size=23, fill=MUTED)
    rounded(draw, (914, 136, 984, 206), 22, "#EFF6FF")
    draw_text(draw, (949, 171), "HT", size=22, fill=BLUE, bold=True, anchor="mm")
    draw.line((0, 224, WIDTH, 224), fill=LINE, width=2)

    rounded(draw, (44, 258, 1036, 370), 16, SURFACE, LINE)
    rounded(draw, (74, 288, 130, 344), 14, BLUE)
    draw_text(draw, (102, 316), "!", size=34, fill=SURFACE, bold=True, anchor="mm")
    draw_text(draw, (156, 280), "Cần hỗ trợ ngay?", size=31, bold=True)
    draw_text(draw, (156, 318), "Chọn dịch vụ, xem giá dự kiến và kết nối nhà cung cấp gần bạn", size=22, fill=MUTED, max_width=560)
    rounded(draw, (790, 294, 1010, 340), 12, GREEN)
    draw_text(draw, (900, 317), "+ Tạo yêu cầu", size=21, fill=SURFACE, bold=True, anchor="mm")

    draw_text(draw, (44, 420), "Dịch vụ phổ biến", size=40, bold=True)
    draw_text(draw, (44, 468), "Đặt nhanh sửa chữa, giao hàng và ăn uống", size=24, fill=MUTED)
    x = 44
    for label, selected in [("Tất cả", True), ("Sửa chữa", False), ("Giao hàng", False), ("Đặt đồ ăn", False)]:
        x = chip(draw, x, 516, label, selected)

    services = [
        ("Sửa điện nước tại nhà", "Thợ xác minh, báo giá trước khi làm", "Từ 120.000đ", "30-45 phút", BLUE, "Phổ biến"),
        ("Giao hàng nội thành", "Xe máy, xe tải nhỏ, giao nhanh trong ngày", "Từ 18.000đ", "15 phút", GREEN, "Nhanh"),
        ("Đặt đồ ăn gần bạn", "Quán cơm, cafe, món Việt, đồ uống", "Ưu đãi 25%", "20-35 phút", WARM, "Deal"),
        ("Vệ sinh máy lạnh", "Kiểm tra gas, vệ sinh dàn lạnh, bảo hành 7 ngày", "Từ 180.000đ", "Trong 2 giờ", "#14B8A6", ""),
    ]
    for (title, subtitle, price, eta, color, badge), (x, y) in zip(services, [(44, 590), (558, 590), (44, 852), (558, 852)]):
        rounded(draw, (x, y, x + 478, y + 238), 16, SURFACE, LINE)
        rounded(draw, (x + 18, y + 18, x + 84, y + 84), 16, f"{color}22")
        draw_text(draw, (x + 51, y + 52), "DV", size=22, fill=color, bold=True, anchor="mm")
        if badge:
            rounded(draw, (x + 288, y + 20, x + 452, y + 54), 17, "#EFF6FF")
            draw_text(draw, (x + 370, y + 38), badge, size=18, fill=BLUE, bold=True, anchor="mm")
        draw_text(draw, (x + 18, y + 100), title, size=27, fill=INK, bold=True, max_width=420)
        draw_text(draw, (x + 18, y + 166), subtitle, size=21, fill=MUTED, max_width=420)
        draw_text(draw, (x + 18, y + 205), price, size=25, fill=BLUE, bold=True)
        draw_text(draw, (x + 452, y + 209), eta, size=20, fill=MUTED, anchor="ra")

    draw_text(draw, (44, 1148), "Nhà cung cấp sẵn sàng", size=36, bold=True)
    providers = [
        ("Tổ thợ An Tâm", "Sửa điện nước · Quận 1, Quận 3", "★ 4.9 · Có mặt 35 phút", BLUE),
        ("ShipNow Local", "Giao hàng nhanh · Nội thành TP.HCM", "★ 4.8 · Nhận đơn 12 phút", GREEN),
        ("Bếp Nhà Gần Đây", "Đặt đồ ăn · Bán kính 3 km", "★ 4.6 · Giao 25 phút", WARM),
    ]
    for index, (name, subtitle, meta, color) in enumerate(providers):
        y = 1210 + index * 112
        rounded(draw, (44, y, 1036, y + 92), 14, SURFACE, LINE)
        rounded(draw, (70, y + 18, 126, y + 74), 14, f"{color}22")
        draw_text(draw, (98, y + 47), "NC", size=18, fill=color, bold=True, anchor="mm")
        verified_suffix = " (xác minh)" if index < 2 else ""
        draw_text(draw, (150, y + 15), f"{name}{verified_suffix}", size=28, bold=True)
        draw_text(draw, (150, y + 49), subtitle, size=22, fill=MUTED)
        draw_text(draw, (150, y + 73), meta, size=20, fill=MUTED)
        rounded(draw, (944, y + 26, 1000, y + 82), 18, "#EFF6FF")
        draw_text(draw, (972, y + 54), "Chat", size=17, fill=BLUE, bold=True, anchor="mm")

    bottom_tabs(draw, "services")
    target = ASSETS / "tab_dich_vu.png"
    image.save(target)
    return target


def make_personal_asset() -> Path:
    image, draw = phone_canvas()
    draw.ellipse((WIDTH // 2 - 84, 180, WIDTH // 2 + 84, 348), fill=GREEN)
    draw_text(draw, (WIDTH // 2, 264), "Q", size=72, fill=SURFACE, bold=True, anchor="mm")
    draw_text(draw, (WIDTH // 2, 384), "Quoc", size=44, fill=INK, bold=True, anchor="mm")
    draw_text(draw, (WIDTH // 2, 432), "quoc@example.com", size=27, fill=MUTED, anchor="mm")
    draw_text(draw, (WIDTH // 2, 478), "Nhấn để chỉnh sửa hồ sơ", size=25, fill=BLUE, bold=True, anchor="mm")

    rounded(draw, (44, 550, 1036, 816), 18, SURFACE, LINE)
    rows = [("Thông báo", "Bật"), ("Quyền riêng tư", ">"), ("Giới thiệu", ">")]
    for index, (label, right) in enumerate(rows):
        y = 550 + index * 88
        if index > 0:
            draw.line((84, y, 1000, y), fill=LINE, width=1)
        draw_text(draw, (88, y + 44), "●", size=28, fill=BLUE, bold=True, anchor="mm")
        draw_text(draw, (132, y + 28), label, size=30, fill=INK, bold=True)
        if right == "Bật":
            rounded(draw, (900, y + 26, 994, y + 62), 18, "#93C5FD")
            draw.ellipse((960, y + 28, 992, y + 60), fill=BLUE)
        else:
            draw_text(draw, (976, y + 44), right, size=42, fill=FAINT, bold=True, anchor="mm")

    rounded(draw, (44, 870, 1036, 934), 16, DANGER)
    draw_text(draw, (540, 903), "Đăng xuất", size=28, fill=SURFACE, bold=True, anchor="mm")

    rounded(draw, (44, 1004, 1036, 1202), 16, SURFACE, LINE)
    draw_text(draw, (82, 1034), "Chỉnh sửa hồ sơ", size=34, bold=True)
    draw_text(draw, (82, 1084), "• Đổi ảnh đại diện", size=27, fill=MUTED)
    draw_text(draw, (82, 1128), "• Cập nhật tên hiển thị", size=27, fill=MUTED)
    draw_text(draw, (82, 1172), "• Email hiển thị ở chế độ chỉ đọc", size=27, fill=MUTED)

    rounded(draw, (44, 1260, 1036, 1394), 16, "#FFF7ED", "#FED7AA")
    draw_text(draw, (82, 1292), "Bảo mật tài khoản", size=32, fill=INK, bold=True)
    draw_text(draw, (82, 1340), "Tin nhắn truyền qua TLS; mã hóa đầu cuối đang được phát triển.", size=25, fill=MUTED, max_width=880)

    bottom_tabs(draw, "personal")
    target = ASSETS / "tab_ca_nhan.png"
    image.save(target)
    return target


def configure_doc(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)

    for style_name, size, color in [
        ("Title", 24, "111827"),
        ("Heading 1", 17, "1D4ED8"),
        ("Heading 2", 14, "111827"),
        ("Heading 3", 12, "111827"),
    ]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    caption = document.styles.add_style("CaptionCustom", WD_STYLE_TYPE.PARAGRAPH)
    caption.font.name = "Arial"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor(102, 112, 133)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, value: str, bold: bool = False, color: str = "111827") -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(value)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(10)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header_value in enumerate(headers):
        cell_text(table.rows[0].cells[index], header_value, bold=True, color="FFFFFF")
        shade_cell(table.rows[0].cells[index], "2563EB")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell_text(cells[index], value)
    document.add_paragraph()


def paragraph(document: Document, value: str = "") -> None:
    document.add_paragraph(value)


def bullets(document: Document, values: Iterable[str]) -> None:
    for value in values:
        document.add_paragraph(value, style="List Bullet")


def numbers(document: Document, values: Iterable[str]) -> None:
    for value in values:
        document.add_paragraph(value, style="List Number")


def add_image(document: Document, path: Path, caption: str, width: float = 2.75) -> None:
    paragraph_obj = document.add_paragraph()
    paragraph_obj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_obj.add_run().add_picture(str(path), width=Inches(width))
    caption_paragraph = document.add_paragraph(style="CaptionCustom")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.add_run(caption)


def build_report(images: dict[str, Path]) -> None:
    document = Document()
    configure_doc(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BÁO CÁO TỔNG QUÁT ỨNG DỤNG KOOLA")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("2563EB")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Chat • Mua sắm • Kết nối doanh nghiệp • Dịch vụ • Cá nhân")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string("667085")

    add_image(document, images["chat"], "Ảnh chụp thực tế tab Trò chuyện trong ứng dụng Koola.", width=2.55)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Ngày lập báo cáo: 22/05/2026").bold = True
    meta.add_run("\nNguồn phân tích: mã nguồn APP_KOOLA tại D:\\Desktop\\APP_KOOLA")

    document.add_page_break()
    document.add_heading("Mục lục", level=1)
    for item in [
        "1. Giới thiệu chung",
        "2. Kiến trúc tổng quan",
        "3. Chức năng chính của ứng dụng",
        "4. Hướng dẫn sử dụng nhanh",
        "5. Nội dung từng tab",
        "6. Luồng hoạt động tiêu biểu",
        "7. Ghi chú triển khai và kết luận",
    ]:
        paragraph(document, item)
    document.add_page_break()

    document.add_heading("1. Giới thiệu chung", level=1)
    paragraph(document, "Koola là ứng dụng di động phục vụ nhắn tin, gọi thoại/gọi video, kết nối doanh nghiệp, mua sắm và đặt dịch vụ trong cùng một trải nghiệm. Ứng dụng được tổ chức dạng monorepo gồm mobile app, backend API và hạ tầng chạy local.")
    add_table(document, ["Thành phần", "Vị trí", "Công nghệ", "Vai trò"], [
        ["Mobile app", "ChatApp/", "React Native 0.76 + TypeScript", "Giao diện người dùng, điều hướng tab, chat, gọi, mua sắm, dịch vụ."],
        ["Backend API", "chat-backend/", "NestJS 11 + MongoDB + Socket.IO", "Xử lý xác thực, hội thoại, tin nhắn, media, doanh nghiệp, WebRTC và realtime."],
        ["Hạ tầng local", "infra-local/", "Docker Compose", "MongoDB, Redis, MinIO và Coturn cho môi trường phát triển."],
    ])
    paragraph(document, "Mục tiêu sản phẩm: tạo một nền tảng giao tiếp và giao dịch nội bộ/địa phương, trong đó người dùng có thể trò chuyện, gọi điện, tìm đối tác, xem sản phẩm, đặt dịch vụ và quản lý hồ sơ cá nhân.")

    document.add_heading("2. Kiến trúc tổng quan", level=1)
    bullets(document, [
        "Mobile sử dụng React Navigation với 5 tab chính: Trò chuyện, Mua sắm, Kết nối, Dịch vụ và Cá nhân.",
        "Backend tổ chức theo module NestJS: auth, users, conversations, messages, media, gateway, webrtc, businesses, notifications và health.",
        "Realtime chat dùng Socket.IO; Redis pub/sub adapter hỗ trợ fanout khi mở rộng nhiều backend instance.",
        "Media sử dụng MinIO thông qua presigned URL, giúp client tải lên/tải xuống trực tiếp thay vì proxy file qua API.",
        "WebRTC dùng namespace riêng cho tín hiệu gọi thoại/gọi video, Coturn hỗ trợ NAT traversal.",
    ])
    add_table(document, ["Luồng", "Mô tả ngắn"], [
        ["Đăng nhập", "Người dùng nhập email/mật khẩu; app lưu token, tự refresh khi access token hết hạn."],
        ["Gửi tin nhắn", "Mobile gửi REST request; backend lưu MongoDB; gateway phát sự kiện realtime tới phòng hội thoại."],
        ["Gửi media", "Mobile xin URL upload, tải file lên MinIO, rồi gửi message chứa media key."],
        ["Gọi thoại/video", "Mobile phát tín hiệu qua WebRTC service; backend quản lý session, timeout và ICE/TURN."],
        ["Kết nối doanh nghiệp", "Mobile gọi businesses API để tìm kiếm, lọc, kết nối và mở hội thoại với chủ doanh nghiệp."],
    ])

    document.add_heading("3. Chức năng chính của ứng dụng", level=1)
    bullets(document, [
        "Xác thực tài khoản: đăng nhập, đăng ký, xác thực OTP email, refresh token và đăng xuất.",
        "Quản lý hội thoại: danh sách chat, tạo nhóm, mở chat 1-1, xem thông tin nhóm/hồ sơ.",
        "Nhắn tin realtime: gửi/nhận tin, trạng thái nhập tin, đánh dấu đã đọc, đồng bộ tin bỏ lỡ.",
        "Tương tác tin nhắn: thả cảm xúc, sao chép, chuyển tiếp, ghim/bỏ ghim, xóa cho tôi hoặc xóa cho mọi người trong giới hạn thời gian.",
        "Media trong chat: gửi ảnh, video và tài liệu; xem ảnh toàn màn hình; phát video trong modal; tải tệp về máy.",
        "Offline: hiển thị trạng thái mất mạng, xếp hàng tin nhắn khi offline và flush queue khi có mạng lại.",
        "Cuộc gọi: gọi thoại/gọi video 1-1, lịch sử cuộc gọi, gọi lại từ nhật ký, xử lý cuộc gọi đến.",
        "Danh bạ/tìm kiếm: tìm người dùng, bắt đầu hội thoại, tìm hội thoại/tin nhắn trên toàn app.",
        "Kết nối doanh nghiệp: duyệt danh sách doanh nghiệp, lọc theo đối tác/nhà cung cấp, tỉnh/thành, danh mục và sắp xếp.",
        "Mua sắm và dịch vụ: màn hình danh mục, sản phẩm/dịch vụ gợi ý, cửa hàng/nhà cung cấp gần người dùng.",
        "Cá nhân: cập nhật avatar/tên hiển thị, bật/tắt thông báo, xem quyền riêng tư, giới thiệu app và đăng xuất.",
    ])

    document.add_heading("4. Hướng dẫn sử dụng nhanh", level=1)
    document.add_heading("4.1 Đăng nhập và đăng ký", level=2)
    numbers(document, [
        "Mở ứng dụng Koola, nhập email và mật khẩu tại màn hình Đăng nhập.",
        "Nếu chưa có tài khoản, chọn Tạo tài khoản, nhập tên hiển thị, email và mật khẩu tối thiểu 6 ký tự.",
        "Nhập mã OTP 6 số được gửi đến email để hoàn tất xác thực.",
        "Sau khi đăng nhập thành công, ứng dụng tự chuyển vào màn hình chính gồm 5 tab.",
    ])
    document.add_heading("4.2 Sử dụng chat", level=2)
    numbers(document, [
        "Vào tab Trò chuyện, chọn hội thoại trong danh sách hoặc dùng nút + để tạo nhóm.",
        "Nhập nội dung vào ô Nhập tin nhắn, bấm Gửi để gửi tin nhắn.",
        "Bấm biểu tượng đính kèm để chọn Ảnh, Video hoặc Tài liệu.",
        "Nhấn giữ một tin nhắn để mở menu thả cảm xúc, sao chép, chuyển tiếp, ghim hoặc xóa.",
        "Bấm biểu tượng gọi thoại/gọi video trong màn hình chat để bắt đầu cuộc gọi 1-1.",
    ])
    document.add_heading("4.3 Tìm kiếm và kết nối", level=2)
    numbers(document, [
        "Dùng thanh tìm kiếm trên header để tìm hội thoại, tin nhắn hoặc doanh nghiệp tùy tab.",
        "Trong tab Kết nối, chọn Tất cả/Đối tác/Nhà cung cấp, sau đó lọc thêm theo tỉnh, danh mục và sắp xếp.",
        "Bấm một doanh nghiệp để xem hồ sơ; chọn Kết nối ngay hoặc Nhắn tin để bắt đầu trao đổi.",
        "Trong tab Cá nhân, nhấn vào phần hồ sơ để đổi avatar hoặc tên hiển thị.",
    ])

    document.add_heading("5. Nội dung từng tab", level=1)
    tab_sections = [
        (
            "5.1 Tab Trò chuyện",
            "Tab Trò chuyện là trung tâm giao tiếp của ứng dụng. Màn hình có header KOOLA, thanh tìm kiếm, nút QR, nút tạo mới và ba tab con: Tin nhắn, Cuộc gọi, Danh bạ.",
            [
                "Tin nhắn: hiển thị danh sách hội thoại, preview tin nhắn cuối, thời gian gần nhất, trạng thái online và số tin chưa đọc.",
                "Cuộc gọi: hiển thị lịch sử gọi thoại/gọi video, trạng thái đã kết thúc, nhỡ, bị từ chối, bận hoặc thất bại; cho phép gọi lại.",
                "Danh bạ: tìm người theo tên/email, chọn người để tạo hoặc mở hội thoại 1-1.",
                "Màn hình chat: gửi tin, gửi media, xem ảnh/video/tệp, thả cảm xúc, chuyển tiếp, ghim, xóa tin, xem trạng thái nhập tin và gọi 1-1.",
            ],
            "chat",
            "Tab Trò chuyện: danh sách hội thoại, tìm kiếm, QR, tạo nhóm và thanh tab chính.",
        ),
        (
            "5.2 Tab Mua sắm",
            "Tab Mua sắm là màn hình commerce mẫu trong app. Nội dung hiện tại gồm tìm kiếm sản phẩm/cửa hàng, giỏ hàng, banner ưu đãi, nhóm hành động nhanh, danh mục, sản phẩm gợi ý và cửa hàng gần người dùng.",
            [
                "Tìm sản phẩm hoặc cửa hàng qua ô tìm kiếm trên đầu màn hình.",
                "Xem các nhóm nhanh: Siêu thị, Ăn uống, Freeship, Deal sốc.",
                "Lọc sản phẩm theo danh mục: Tất cả, Tạp hóa, Đồ ăn, Điện tử, Nhà cửa, Làm đẹp.",
                "Mỗi sản phẩm có giá, cửa hàng, điểm đánh giá, số đã bán, badge khuyến mãi và nút thêm vào giỏ.",
            ],
            "shopping",
            "Ảnh minh họa tab Mua sắm dựa trên màn hình ShoppingHomeScreen.",
        ),
        (
            "5.3 Tab Kết nối",
            "Tab Kết nối phục vụ tìm và kết nối doanh nghiệp. Đây là khu vực cho đối tác, nhà cung cấp, doanh nghiệp địa phương hoặc đơn vị dịch vụ B2B.",
            [
                "Header có tìm kiếm doanh nghiệp, QR và nút đăng ký doanh nghiệp mới.",
                "Bộ lọc quan hệ gồm Tất cả, Đối tác và Nhà cung cấp.",
                "Filter bar hỗ trợ chọn tỉnh/thành, sắp xếp và lọc theo lĩnh vực như Logistics, Nội địa, Nguyên liệu, Bao bì, Sản xuất, Thực phẩm, Công nghệ, Tài chính, Bất động sản, Bán lẻ, Y tế, Giáo dục.",
                "Business card hiển thị tên doanh nghiệp, lĩnh vực, mô tả, trạng thái xác minh/kết nối và nút nhắn tin.",
            ],
            "connect",
            "Ảnh minh họa tab Kết nối: lọc doanh nghiệp, danh sách business card và thao tác nhắn tin.",
        ),
        (
            "5.4 Tab Dịch vụ",
            "Tab Dịch vụ gom các nhu cầu hằng ngày như sửa chữa, giao hàng, đặt đồ ăn và dịch vụ nhà cửa. Giao diện hiện tại là màn hình danh mục/dữ liệu mẫu để trình bày trải nghiệm đặt dịch vụ.",
            [
                "Header có tiêu đề Dịch vụ, thanh tìm kiếm và nút trung tâm hỗ trợ.",
                "Banner Cần hỗ trợ ngay cho phép tạo yêu cầu dịch vụ.",
                "Danh mục dịch vụ gồm Tất cả, Sửa chữa, Giao hàng, Đặt đồ ăn và Nhà cửa.",
                "Mỗi thẻ dịch vụ hiển thị tên, mô tả, giá dự kiến, thời gian đáp ứng, điểm đánh giá và badge nổi bật.",
            ],
            "services",
            "Ảnh minh họa tab Dịch vụ: dịch vụ phổ biến và nhà cung cấp sẵn sàng.",
        ),
        (
            "5.5 Tab Cá nhân",
            "Tab Cá nhân dùng để quản lý tài khoản và các thiết lập cơ bản của người dùng.",
            [
                "Hiển thị avatar, tên người dùng, email và lối vào chỉnh sửa hồ sơ.",
                "Chỉnh sửa hồ sơ cho phép đổi ảnh đại diện và tên hiển thị; email được hiển thị dạng chỉ đọc.",
                "Cài đặt thông báo có công tắc bật/tắt và gọi API cập nhật user settings.",
                "Quyền riêng tư hiển thị thông tin bảo mật dữ liệu và trạng thái mã hóa đầu cuối đang phát triển.",
                "Giới thiệu hiển thị phiên bản ứng dụng và thông tin công nghệ.",
                "Nút Đăng xuất kết thúc phiên đăng nhập và ngắt kết nối socket/WebRTC.",
            ],
            "personal",
            "Ảnh minh họa tab Cá nhân: hồ sơ, cài đặt thông báo, quyền riêng tư và đăng xuất.",
        ),
    ]
    for heading, intro, items, image_key, caption in tab_sections:
        document.add_heading(heading, level=2)
        paragraph(document, intro)
        bullets(document, items)
        add_image(document, images[image_key], caption, width=2.7)

    document.add_heading("6. Luồng hoạt động tiêu biểu", level=1)
    document.add_heading("6.1 Luồng gửi tin nhắn", level=2)
    numbers(document, [
        "Người dùng nhập tin trong ChatScreen và bấm Gửi.",
        "Mobile gọi API gửi tin nhắn hoặc đưa vào OfflineQueue nếu mất mạng.",
        "Backend kiểm tra người gửi có thuộc hội thoại hay không, lưu message vào MongoDB và cập nhật lastMessage của conversation.",
        "Gateway phát sự kiện new_message tới room hội thoại qua Socket.IO/Redis adapter.",
        "Các client đang online cập nhật danh sách chat và nội dung hội thoại theo realtime.",
    ])
    document.add_heading("6.2 Luồng gửi media", level=2)
    numbers(document, [
        "Người dùng chọn Ảnh, Video hoặc Tài liệu từ menu đính kèm.",
        "Mobile kiểm tra kích thước/định dạng; video có thể được nén trước khi upload.",
        "Mobile xin presigned upload URL từ backend media API.",
        "File được upload trực tiếp lên MinIO; message lưu mediaKey, MIME type, kích thước và metadata liên quan.",
        "Khi xem lại, mobile xin download URL hoặc dùng cache local để hiển thị nhanh hơn.",
    ])
    document.add_heading("6.3 Luồng gọi thoại/gọi video", level=2)
    numbers(document, [
        "Người dùng bấm gọi thoại hoặc gọi video trong màn hình chat.",
        "Mobile gửi sự kiện khởi tạo cuộc gọi qua WebRTC service.",
        "Backend tạo call session, cấp ICE/TURN server và gửi incoming call tới người nhận.",
        "Hai phía trao đổi offer/answer/ICE candidate để thiết lập kết nối WebRTC.",
        "Kết thúc, nhỡ, từ chối hoặc lỗi cuộc gọi được ghi vào call logs để hiển thị trong tab Cuộc gọi.",
    ])

    document.add_heading("7. Ghi chú triển khai và kết luận", level=1)
    bullets(document, [
        "Báo cáo này được tổng hợp từ cấu trúc code hiện tại của APP_KOOLA và các tài liệu kỹ thuật trong thư mục docs/.",
        "Ảnh tab Trò chuyện là ảnh chụp có sẵn trong repo. Ảnh các tab Mua sắm, Kết nối, Dịch vụ và Cá nhân là mockup minh họa được dựng theo code màn hình tương ứng để phục vụ báo cáo.",
        "Một số màn hình commerce/dịch vụ đang sử dụng dữ liệu mock, phù hợp cho demo giao diện và luồng sử dụng ban đầu.",
        "Các chức năng cốt lõi của app nằm ở chat realtime, media, offline queue, call WebRTC, tìm kiếm và kết nối doanh nghiệp.",
    ])
    paragraph(document, "Kết luận: Koola có nền tảng kỹ thuật tương đối đầy đủ cho một ứng dụng chat mở rộng, không chỉ dừng ở nhắn tin mà còn hướng tới hệ sinh thái kết nối, mua sắm và dịch vụ. Với backend NestJS, MongoDB, Socket.IO, Redis, MinIO và Coturn, hệ thống đã có các thành phần cần thiết để phát triển tiếp theo hướng production.")

    document.save(OUT)


def main() -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)
    images = {
        "chat": save_chat_asset(),
        "shopping": make_shopping_asset(),
        "connect": make_connect_asset(),
        "services": make_services_asset(),
        "personal": make_personal_asset(),
    }
    build_report(images)
    print(OUT)
    for image_path in images.values():
        print(image_path)


if __name__ == "__main__":
    main()
